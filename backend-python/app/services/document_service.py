"""
Document processing service for PDF and DOCX files.
Handles text extraction, OCR, table extraction, and metadata parsing.
"""
import logging
from typing import Dict, List, Optional, Any, BinaryIO
from pathlib import Path
from datetime import datetime
import io

# PDF Processing
import pypdf  # Replaced PyPDF2 - maintained fork with security fixes
import pdfplumber

# OCR
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR not available - pytesseract or Pillow not installed")

# DOCX Processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("DOCX processing not available - python-docx not installed")

logger = logging.getLogger(__name__)


class DocumentProcessingError(Exception):
    """Custom exception for document processing errors."""
    pass


class DocumentService:
    """Service for processing various document formats."""

    @staticmethod
    def extract_pdf_text(file_path: str | Path | BinaryIO) -> Dict[str, Any]:
        """
        Extract text from PDF using pypdf.
        Fast and reliable for text-based PDFs.

        Args:
            file_path: Path to PDF file or file-like object

        Returns:
            Dict with extracted text, page count, and metadata
        """
        try:
            # Handle both file paths and file-like objects
            if isinstance(file_path, (str, Path)):
                pdf_file = open(file_path, 'rb')
                should_close = True
            else:
                pdf_file = file_path
                should_close = False

            try:
                reader = pypdf.PdfReader(pdf_file)

                # Extract metadata
                metadata = {
                    'title': reader.metadata.title if reader.metadata else None,
                    'author': reader.metadata.author if reader.metadata else None,
                    'subject': reader.metadata.subject if reader.metadata else None,
                    'creator': reader.metadata.creator if reader.metadata else None,
                    'producer': reader.metadata.producer if reader.metadata else None,
                    'creation_date': reader.metadata.creation_date if reader.metadata else None,
                }

                # Extract text from all pages
                pages = []
                full_text = []

                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    pages.append({
                        'page_number': i + 1,
                        'text': page_text,
                        'char_count': len(page_text)
                    })
                    full_text.append(page_text)

                return {
                    'success': True,
                    'text': '\n\n'.join(full_text),
                    'page_count': len(reader.pages),
                    'pages': pages,
                    'metadata': metadata,
                    'processing_method': 'PyPDF2',
                }
            finally:
                if should_close:
                    pdf_file.close()

        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise DocumentProcessingError(f"Failed to extract PDF text: {str(e)}")

    @staticmethod
    def extract_pdf_tables(file_path: str | Path | BinaryIO) -> Dict[str, Any]:
        """
        Extract tables from PDF using pdfplumber.
        Better for structured data and tables.

        Args:
            file_path: Path to PDF file or file-like object

        Returns:
            Dict with extracted text, tables, and page info
        """
        try:
            # pdfplumber can handle both paths and file-like objects
            with pdfplumber.open(file_path) as pdf:
                pages_data = []
                all_tables = []
                full_text = []

                for i, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text() or ""
                    full_text.append(page_text)

                    # Extract tables
                    tables = page.extract_tables()
                    page_tables = []

                    for table_idx, table in enumerate(tables):
                        if table:
                            # Convert table to dict format
                            headers = table[0] if table else []
                            rows = table[1:] if len(table) > 1 else []

                            table_data = {
                                'table_number': table_idx + 1,
                                'headers': headers,
                                'rows': rows,
                                'row_count': len(rows),
                                'column_count': len(headers)
                            }
                            page_tables.append(table_data)
                            all_tables.append({
                                'page': i + 1,
                                **table_data
                            })

                    pages_data.append({
                        'page_number': i + 1,
                        'text': page_text,
                        'tables': page_tables,
                        'table_count': len(page_tables)
                    })

                return {
                    'success': True,
                    'text': '\n\n'.join(full_text),
                    'page_count': len(pdf.pages),
                    'pages': pages_data,
                    'tables': all_tables,
                    'total_tables': len(all_tables),
                    'processing_method': 'pdfplumber',
                }

        except Exception as e:
            logger.error(f"Error extracting PDF tables: {e}")
            raise DocumentProcessingError(f"Failed to extract PDF tables: {str(e)}")

    @staticmethod
    def ocr_pdf(file_path: str | Path, language: str = 'eng') -> Dict[str, Any]:
        """
        Perform OCR on scanned PDF documents.
        Uses pytesseract for optical character recognition.

        Args:
            file_path: Path to PDF file
            language: Tesseract language code (default: 'eng')

        Returns:
            Dict with OCR-extracted text and confidence scores
        """
        if not OCR_AVAILABLE:
            raise DocumentProcessingError("OCR not available - install pytesseract and Pillow")

        try:
            # Convert PDF pages to images and OCR
            from pdf2image import convert_from_path

            pages = convert_from_path(file_path, dpi=300)

            pages_data = []
            full_text = []

            for i, page_image in enumerate(pages):
                # Perform OCR on page
                ocr_data = pytesseract.image_to_data(
                    page_image,
                    lang=language,
                    output_type=pytesseract.Output.DICT
                )

                # Extract text
                page_text = pytesseract.image_to_string(page_image, lang=language)
                full_text.append(page_text)

                # Calculate average confidence
                confidences = [int(conf) for conf in ocr_data['conf'] if conf != '-1']
                avg_confidence = sum(confidences) / len(confidences) if confidences else 0

                pages_data.append({
                    'page_number': i + 1,
                    'text': page_text,
                    'confidence': round(avg_confidence, 2),
                    'word_count': len(page_text.split())
                })

            return {
                'success': True,
                'text': '\n\n'.join(full_text),
                'page_count': len(pages),
                'pages': pages_data,
                'processing_method': 'OCR (Tesseract)',
                'language': language,
            }

        except ImportError:
            raise DocumentProcessingError("pdf2image not installed - required for OCR")
        except Exception as e:
            logger.error(f"Error performing OCR: {e}")
            raise DocumentProcessingError(f"Failed to perform OCR: {str(e)}")

    @staticmethod
    def extract_docx_text(file_path: str | Path | BinaryIO) -> Dict[str, Any]:
        """
        Extract text and structure from DOCX files.

        Args:
            file_path: Path to DOCX file or file-like object

        Returns:
            Dict with extracted text, paragraphs, tables, and metadata
        """
        if not DOCX_AVAILABLE:
            raise DocumentProcessingError("DOCX processing not available - install python-docx")

        try:
            doc = Document(file_path)

            # Extract metadata
            core_properties = doc.core_properties
            metadata = {
                'title': core_properties.title,
                'author': core_properties.author,
                'subject': core_properties.subject,
                'created': core_properties.created,
                'modified': core_properties.modified,
                'last_modified_by': core_properties.last_modified_by,
            }

            # Extract paragraphs
            paragraphs = []
            full_text = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        'text': para.text,
                        'style': para.style.name,
                    })
                    full_text.append(para.text)

            # Extract tables
            tables_data = []
            for table_idx, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    rows.append([cell.text for cell in row.cells])

                tables_data.append({
                    'table_number': table_idx + 1,
                    'rows': rows,
                    'row_count': len(rows),
                    'column_count': len(rows[0]) if rows else 0
                })

            return {
                'success': True,
                'text': '\n\n'.join(full_text),
                'paragraphs': paragraphs,
                'paragraph_count': len(paragraphs),
                'tables': tables_data,
                'table_count': len(tables_data),
                'metadata': metadata,
                'processing_method': 'python-docx',
            }

        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise DocumentProcessingError(f"Failed to extract DOCX text: {str(e)}")

    @staticmethod
    def process_document(
        file_path: str | Path | BinaryIO,
        file_type: Optional[str] = None,
        extract_tables: bool = True,
        use_ocr: bool = False,
        ocr_language: str = 'eng'
    ) -> Dict[str, Any]:
        """
        Intelligently process any document format.
        Automatically detects file type and uses appropriate processing method.

        Args:
            file_path: Path to file or file-like object
            file_type: Optional file type hint ('pdf', 'docx')
            extract_tables: Whether to extract tables from PDFs (slower but more detailed)
            use_ocr: Whether to use OCR for scanned PDFs
            ocr_language: Language for OCR processing

        Returns:
            Dict with processed document data
        """
        start_time = datetime.now()

        try:
            # Detect file type if not provided
            if file_type is None:
                if isinstance(file_path, (str, Path)):
                    file_type = Path(file_path).suffix.lower().lstrip('.')
                else:
                    raise DocumentProcessingError("file_type must be provided for file-like objects")

            result = None

            # Process based on file type
            if file_type == 'pdf':
                if use_ocr and isinstance(file_path, (str, Path)):
                    result = DocumentService.ocr_pdf(file_path, language=ocr_language)
                elif extract_tables:
                    result = DocumentService.extract_pdf_tables(file_path)
                else:
                    result = DocumentService.extract_pdf_text(file_path)

            elif file_type in ['docx', 'doc']:
                result = DocumentService.extract_docx_text(file_path)

            else:
                raise DocumentProcessingError(f"Unsupported file type: {file_type}")

            # Add processing time
            processing_time = (datetime.now() - start_time).total_seconds()
            result['processing_time_seconds'] = round(processing_time, 3)
            result['file_type'] = file_type

            logger.info(f"Successfully processed {file_type} document in {processing_time:.2f}s")

            return result

        except Exception as e:
            logger.error(f"Error processing document: {e}")
            raise DocumentProcessingError(f"Failed to process document: {str(e)}")

    @staticmethod
    def extract_key_information(text: str) -> Dict[str, Any]:
        """
        Extract key information from document text.
        Useful for contact extraction, date parsing, etc.

        Args:
            text: Document text to analyze

        Returns:
            Dict with extracted information
        """
        import re

        # Email extraction
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)

        # Phone number extraction (various formats)
        phone_pattern = r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, text)

        # URL extraction
        url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
        urls = re.findall(url_pattern, text)

        # Date extraction (common formats)
        date_pattern = r'\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}\b'
        dates = re.findall(date_pattern, text, re.IGNORECASE)

        return {
            'emails': list(set(emails)),
            'phone_numbers': list(set(phones)),
            'urls': list(set(urls)),
            'dates': list(set(dates)),
            'word_count': len(text.split()),
            'character_count': len(text),
        }
